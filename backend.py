#!/usr/bin/env python3
import os
import re
import json
import hmac
import base64
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import requests
from datetime import datetime, timezone
from decimal import Decimal
import boto3
from boto3.dynamodb.conditions import Key

try:
    from cryptography.fernet import Fernet
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

app = Flask(__name__)
CORS(app)

OPENAI_API_KEY = os.environ.get('OPENAI_API_KEY', '')
TRANSCRIPTION_MODEL = os.environ.get('OPENAI_TRANSCRIPTION_MODEL', 'gpt-4o-mini-transcribe')
CAREEROPS_USERNAME = os.environ.get('CAREEROPS_USERNAME', '')
CAREEROPS_PASSWORD = os.environ.get('CAREEROPS_PASSWORD', '')
AWS_REGION = os.environ.get('AWS_REGION', 'us-east-1')

# Initialize DynamoDB using the standard AWS credential provider chain.
try:
    dynamodb = boto3.resource('dynamodb', region_name=AWS_REGION)
    applications_table = dynamodb.Table('applications')
    runbook_table = dynamodb.Table('runbook')
    llm_providers_table = dynamodb.Table('llm-providers')
    DYNAMODB_AVAILABLE = True
except Exception as e:
    print(f"⚠️  DynamoDB initialization failed: {e}")
    DYNAMODB_AVAILABLE = False

# ---------- Encryption Setup ----------
# Fernet symmetric encryption for API keys at rest.
# Key source: CAREEROPS_ENCRYPTION_KEY env var (base64-encoded 32-byte key).
# If not set, auto-generates one and prints it (you should persist it).
ENCRYPTION_KEY = os.environ.get('CAREEROPS_ENCRYPTION_KEY', '')
if ENCRYPTION_KEY and CRYPTO_AVAILABLE:
    try:
        fernet = Fernet(ENCRYPTION_KEY.encode())
    except Exception as e:
        print(f"⚠️  Invalid CAREEROPS_ENCRYPTION_KEY: {e}")
        fernet = None
elif CRYPTO_AVAILABLE:
    # Auto-generate for development convenience (NOT for production)
    generated_key = Fernet.generate_key().decode()
    fernet = Fernet(generated_key.encode())
    print(f"⚠️  No CAREEROPS_ENCRYPTION_KEY set. Auto-generated (non-persistent):")
    print(f"   export CAREEROPS_ENCRYPTION_KEY={generated_key}")
    print(f"   Add this to your environment to persist keys across restarts.\n")
else:
    fernet = None


def encrypt_value(plaintext):
    """Encrypt a string value. Returns base64-encoded ciphertext."""
    if not fernet:
        raise RuntimeError("Encryption not available. Install cryptography package and set CAREEROPS_ENCRYPTION_KEY.")
    return fernet.encrypt(plaintext.encode()).decode()


def decrypt_value(ciphertext):
    """Decrypt a base64-encoded ciphertext. Returns plaintext string."""
    if not fernet:
        raise RuntimeError("Decryption not available. Install cryptography package and set CAREEROPS_ENCRYPTION_KEY.")
    return fernet.decrypt(ciphertext.encode()).decode()


def mask_key(key):
    """Return a masked version of an API key for display (e.g., sk-...abc123)."""
    if not key or len(key) < 8:
        return '••••••••'
    return key[:3] + '•' * (len(key) - 7) + key[-4:]

# Initialize S3 client for CV file storage.
S3_BUCKET = 'careerops-589535355002'
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5 MB
ALLOWED_EXTENSIONS = {'.pdf', '.docx'}
ALLOWED_CONTENT_TYPES = {
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
}

try:
    s3_client = boto3.client('s3', region_name=AWS_REGION)
    S3_AVAILABLE = True
except Exception as e:
    print(f"⚠️  S3 initialization failed: {e}")
    S3_AVAILABLE = False

# ========== CV COMPARISON ENGINE ==========


class ParseError(Exception):
    """Raised when CV markdown cannot be parsed into the CV_Category_Set."""
    pass


class CVParser:
    """Parses CV markdown into the fixed CV_Category_Set structure."""

    CATEGORY_HEADERS = {
        'PROFESSIONAL SUMMARY': 'summary',
        'CORE SKILLS': 'core_skills',
        'PROFESSIONAL EXPERIENCE': 'professional_experience',
        'CERTIFICATIONS': 'certifications',
        'EDUCATION': 'education',
        'COMMUNITY LEADERSHIP & ENGAGEMENT': 'community',
        'LANGUAGES': 'languages',
    }

    # The fixed ordering of top-level categories
    CATEGORY_ORDER = [
        'heading', 'subheading', 'contact', 'summary', 'core_skills',
        'professional_experience', 'certifications', 'education',
        'community', 'languages',
    ]

    def parse(self, markdown: str) -> dict:
        """Parse CV markdown into category structure.

        Args:
            markdown: Raw CV markdown string.

        Returns:
            Dict with keys for each category containing content strings.
            Professional Experience includes a 'roles' list of sub-categories.

        Raises:
            ParseError: If input is empty or has no recognizable headers.
        """
        if not markdown or not markdown.strip():
            raise ParseError("Input CV markdown is empty.")

        lines = markdown.split('\n')
        categories = {}
        current_category = None
        found_h1 = False
        found_h2 = False
        found_subheading = False
        found_contact = False
        current_lines = []

        def _flush_lines(cat_key, line_list):
            """Save accumulated lines to the category."""
            if cat_key and line_list:
                # Strip trailing empty lines
                while line_list and line_list[-1].strip() == '':
                    line_list.pop()
                if not line_list:
                    return
                content = '\n'.join(line_list)
                if cat_key == '_after_pe':
                    # Unmapped content after Professional Experience —
                    # append to last PE role's content (nearest preceding category)
                    pe = categories.get('professional_experience')
                    if pe and pe.get('roles'):
                        last_role = pe['roles'][-1]
                        if last_role['content']:
                            last_role['content'] = last_role['content'] + '\n' + content
                        else:
                            last_role['content'] = content
                    # If no PE roles exist, silently discard (shouldn't happen)
                elif cat_key in categories:
                    categories[cat_key] = categories[cat_key] + '\n' + content
                else:
                    categories[cat_key] = content

        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Detect H1 header (heading)
            if stripped.startswith('# ') and not stripped.startswith('## '):
                if not found_h1:
                    _flush_lines(current_category, current_lines)
                    current_lines = []
                    found_h1 = True
                    categories['heading'] = stripped
                    current_category = 'heading'
                    i += 1
                    continue
                else:
                    # Additional H1 — treat as content of current category
                    current_lines.append(line)
                    i += 1
                    continue

            # Detect subheading: first bold line (**...**) after H1, before any H2
            if (found_h1 and not found_subheading and not found_h2
                    and stripped.startswith('**') and stripped.endswith('**')):
                _flush_lines(current_category, current_lines)
                current_lines = []
                found_subheading = True
                categories['subheading'] = stripped
                current_category = 'subheading'
                i += 1
                continue

            # Detect contact line: contains email-like pattern and pipe separators, after subheading/heading, before H2
            if (found_h1 and not found_contact and not found_h2
                    and not stripped.startswith('#')
                    and '|' in stripped and '@' in stripped):
                _flush_lines(current_category, current_lines)
                current_lines = []
                found_contact = True
                categories['contact'] = stripped
                current_category = 'contact'
                i += 1
                continue

            # Detect H2 headers (## SECTION NAME)
            if stripped.startswith('## '):
                found_h2 = True
                header_text = stripped[3:].strip()
                # Look up in CATEGORY_HEADERS
                cat_key = self.CATEGORY_HEADERS.get(header_text.upper())
                if cat_key:
                    _flush_lines(current_category, current_lines)
                    current_lines = []
                    if cat_key == 'professional_experience':
                        # Parse the professional experience section specially
                        pe_result, new_i = self._parse_professional_experience(lines, i)
                        categories['professional_experience'] = pe_result
                        # Set current_category to a marker so unmapped content
                        # after PE gets appended to the last role
                        current_category = '_after_pe'
                        current_lines = []
                        i = new_i
                        continue
                    else:
                        current_category = cat_key
                        current_lines = [stripped]
                        i += 1
                        continue
                else:
                    # Unrecognized H2 header — append to current category
                    current_lines.append(line)
                    i += 1
                    continue

            # Regular line — append to current category accumulator
            current_lines.append(line)
            i += 1

        # Flush remaining lines
        _flush_lines(current_category, current_lines)

        # Validate that we found at least some recognizable structure
        if not found_h1 and not found_h2:
            raise ParseError(
                "No recognizable header structure found. "
                "Expected at least an H1 heading (# ) or H2 section headers (## )."
            )

        return categories

    def _parse_professional_experience(self, lines, start_idx):
        """Parse the Professional Experience section into roles.

        Args:
            lines: All lines of the CV markdown.
            start_idx: Index of the ## PROFESSIONAL EXPERIENCE line.

        Returns:
            Tuple of (pe_dict, next_index) where pe_dict has 'header' and 'roles'.
        """
        header_line = lines[start_idx].strip()
        result = {
            'header': header_line,
            'roles': [],
        }

        i = start_idx + 1
        current_role = None
        role_content_lines = []
        role_idx = 0

        def _flush_role():
            nonlocal current_role, role_content_lines, role_idx
            if current_role:
                current_role['content'] = '\n'.join(role_content_lines).strip()
                result['roles'].append(current_role)
                role_content_lines = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Check if we've hit another H2 section (end of professional experience)
            if stripped.startswith('## '):
                _flush_role()
                return result, i

            # Check for H3 header (role title)
            if stripped.startswith('### '):
                _flush_role()
                current_role = {
                    'key': f'role_{role_idx}',
                    'title': stripped,
                    'metadata': '',
                    'content': '',
                }
                role_content_lines = []
                role_idx += 1

                # Look for metadata line (bold line immediately after H3)
                # Skip empty lines between H3 and metadata
                j = i + 1
                while j < len(lines) and lines[j].strip() == '':
                    j += 1
                if j < len(lines):
                    meta_stripped = lines[j].strip()
                    if meta_stripped.startswith('**') and meta_stripped.endswith('**'):
                        current_role['metadata'] = meta_stripped
                        i = j + 1
                        continue

                i += 1
                continue

            # Regular content within a role or pre-role content
            if current_role is not None:
                role_content_lines.append(line)
            else:
                # Content between the ## header and first ### role
                # Append to nearest preceding category (the PE header context)
                # We'll just skip empty lines here
                pass

            i += 1

        # Flush the last role at end of file
        _flush_role()
        return result, i

    def serialize(self, categories: dict) -> str:
        """Serialize parsed categories back to markdown.

        Args:
            categories: Parsed category dictionary.

        Returns:
            Markdown string equivalent to the original input.
        """
        parts = []

        for cat_key in self.CATEGORY_ORDER:
            if cat_key not in categories:
                continue

            value = categories[cat_key]

            if cat_key == 'professional_experience':
                # Serialize PE with its header and roles
                if isinstance(value, dict):
                    parts.append(value.get('header', '## PROFESSIONAL EXPERIENCE'))
                    parts.append('')  # blank line after header
                    for role in value.get('roles', []):
                        parts.append(role.get('title', ''))
                        if role.get('metadata'):
                            parts.append(role['metadata'])
                        parts.append('')  # blank line before content
                        if role.get('content'):
                            parts.append(role['content'])
                        parts.append('')  # blank line after role
                else:
                    # Fallback: PE stored as plain string
                    parts.append(value)
                    parts.append('')
            elif cat_key in ('heading', 'subheading', 'contact'):
                # These are single-line categories
                parts.append(value)
                parts.append('')
            else:
                # Section categories (summary, core_skills, etc.)
                parts.append(value)
                parts.append('')

        # Join and clean up excessive blank lines
        result = '\n'.join(parts)
        # Normalize multiple blank lines to at most two newlines
        result = re.sub(r'\n{3,}', '\n\n', result)
        return result.strip() + '\n'


def dynamodb_safe(value):
    """Convert JSON payloads into DynamoDB-compatible values."""
    return json.loads(json.dumps(value), parse_float=Decimal)

def json_safe(value):
    if isinstance(value, list):
        return [json_safe(item) for item in value]
    if isinstance(value, dict):
        return {key: json_safe(item) for key, item in value.items()}
    if isinstance(value, Decimal):
        return int(value) if value % 1 == 0 else float(value)
    return value

def validate_cv_file(file):
    """Validate an uploaded CV file. Returns (None) on success or (error_message) on failure."""
    if not file or not file.filename:
        return 'No file provided'

    # Check extension
    filename = os.path.basename(file.filename or '')
    filename = re.sub(r'[^\w\-.]', '_', filename)
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        return f'Invalid file type. Allowed formats: {", ".join(sorted(ALLOWED_EXTENSIONS))}'

    # Check Content-Type matches extension
    expected_content_type = ALLOWED_CONTENT_TYPES.get(ext)
    if file.content_type != expected_content_type:
        return f'Content-Type mismatch. Expected {expected_content_type} for {ext} file'

    # Check file size (read content to verify)
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)

    if size == 0:
        return 'File is empty'

    if size > MAX_FILE_SIZE:
        return f'File too large. Maximum size is 5 MB'

    return None


class ComparisonError(Exception):
    """Raised when comparison inputs are invalid or incomplete."""
    pass


class ComparisonEngine:
    """Compares original and ATS-suggested CV categories."""

    # Required top-level categories (excluding professional_experience which is handled specially)
    _SIMPLE_CATEGORIES = [
        'heading', 'subheading', 'contact', 'summary', 'core_skills',
        'certifications', 'education', 'community', 'languages',
    ]

    def compare(self, original: dict, suggested: dict) -> dict:
        """Produce comparison object with change status per category.

        Args:
            original: Parsed categories from original CV (output of CVParser.parse()).
            suggested: Parsed categories from ATS-suggested CV (output of CVParser.parse()).

        Returns:
            Dict with comparison entries per category in CV_Category_Set order.
            Each entry contains:
            - original_content: str
            - suggested_content: str
            - changed: bool
            - score_impact: int (defaults to 0)

            Professional Experience roles use keys like
            "professional_experience.role_0", "professional_experience.role_1", etc.

        Raises:
            ComparisonError: If either input is invalid/incomplete.
        """
        self._validate_input(original, "original")
        self._validate_input(suggested, "suggested")

        result = {}

        for cat_key in CVParser.CATEGORY_ORDER:
            if cat_key == 'professional_experience':
                # Handle professional experience role-by-role
                self._compare_professional_experience(original, suggested, result)
            else:
                orig_content = original.get(cat_key, '')
                sugg_content = suggested.get(cat_key, '')
                changed = orig_content.strip() != sugg_content.strip()
                result[cat_key] = {
                    'original_content': orig_content,
                    'suggested_content': sugg_content if changed else orig_content,
                    'changed': changed,
                    'score_impact': 0,
                }

        return result

    def _validate_input(self, data: dict, label: str) -> None:
        """Validate that a parsed CV dict is valid and complete.

        Raises ComparisonError if the input is not a dict or is missing
        required category keys.
        """
        if not isinstance(data, dict):
            raise ComparisonError(
                f"The {label} input is not a valid parsed CV dictionary."
            )

        # Must have at least heading and one H2 section to be considered valid
        if 'heading' not in data:
            raise ComparisonError(
                f"The {label} input is missing required 'heading' category."
            )

        # Check that professional_experience, if present, has the expected structure
        pe = data.get('professional_experience')
        if pe is not None and not isinstance(pe, dict):
            raise ComparisonError(
                f"The {label} input has invalid 'professional_experience' structure. "
                "Expected a dict with 'header' and 'roles' keys."
            )
        if isinstance(pe, dict) and 'roles' not in pe:
            raise ComparisonError(
                f"The {label} input 'professional_experience' is missing 'roles' key."
            )

    def _compare_professional_experience(self, original: dict, suggested: dict, result: dict) -> None:
        """Compare Professional Experience header and roles individually.

        Adds a 'professional_experience' entry for the header, and entries like
        "professional_experience.role_0", "professional_experience.role_1" for roles.
        """
        orig_pe = original.get('professional_experience', {})
        sugg_pe = suggested.get('professional_experience', {})

        # Compare the PE header itself
        orig_header = orig_pe.get('header', '') if isinstance(orig_pe, dict) else ''
        sugg_header = sugg_pe.get('header', '') if isinstance(sugg_pe, dict) else ''
        header_changed = orig_header.strip() != sugg_header.strip()
        result['professional_experience'] = {
            'original_content': orig_header,
            'suggested_content': sugg_header if header_changed else orig_header,
            'changed': header_changed,
            'score_impact': 0,
        }

        orig_roles = orig_pe.get('roles', []) if isinstance(orig_pe, dict) else []
        sugg_roles = sugg_pe.get('roles', []) if isinstance(sugg_pe, dict) else []

        # Compare role by role based on position
        max_roles = max(len(orig_roles), len(sugg_roles))

        for i in range(max_roles):
            key = f'professional_experience.role_{i}'

            orig_role = orig_roles[i] if i < len(orig_roles) else None
            sugg_role = sugg_roles[i] if i < len(sugg_roles) else None

            orig_content = self._serialize_role(orig_role) if orig_role else ''
            sugg_content = self._serialize_role(sugg_role) if sugg_role else ''

            changed = orig_content.strip() != sugg_content.strip()
            result[key] = {
                'original_content': orig_content,
                'suggested_content': sugg_content if changed else orig_content,
                'changed': changed,
                'score_impact': 0,
            }

    def _serialize_role(self, role: dict) -> str:
        """Serialize a single role dict to its markdown content string."""
        parts = []
        if role.get('title'):
            parts.append(role['title'])
        if role.get('metadata'):
            parts.append(role['metadata'])
        if role.get('content'):
            parts.append(role['content'])
        return '\n'.join(parts)


# ========== CV ATS COMPARISON UTILITIES ==========


def compute_ats_score(baseline: int, impacts: dict, approvals: dict) -> dict:
    """Compute current and max ATS scores from baseline, per-category impacts, and approval states.

    Args:
        baseline: Base ATS score (0-100).
        impacts: Dict mapping category keys to signed integer score impacts.
        approvals: Dict mapping category keys to boolean (True = approved).

    Returns:
        Dict with 'current_score' and 'max_score', both clamped to [0, 100].
    """
    approved_sum = sum(
        impacts[key] for key in impacts if approvals.get(key, False)
    )
    total_sum = sum(impacts.values())

    current_score = max(0, min(100, baseline + approved_sum))
    max_score = max(0, min(100, baseline + total_sum))

    return {"current_score": current_score, "max_score": max_score}


def truncate_text(text: str, max_len: int) -> str:
    """Truncate text for display purposes.

    If len(text) <= max_len, returns the original text unchanged.
    Otherwise, returns the first max_len characters followed by "…" (ellipsis).

    Args:
        text: The input string.
        max_len: Maximum number of characters before truncation.

    Returns:
        Original text if short enough, or truncated text with "…" appended.
    """
    if len(text) <= max_len:
        return text
    return text[:max_len] + "…"


# ========== CV COMPARISON & ASSEMBLY ==========


class AssemblyError(Exception):
    """Raised when CV assembly fails due to invalid inputs."""
    pass


class FinalAssembler:
    """Assembles a final CV markdown from comparison data and user approvals."""

    def assemble(self, comparison: dict, approvals: dict) -> str:
        """Assemble final CV markdown from comparison entries and approval decisions.

        Args:
            comparison: Dict of comparison entries (from ComparisonEngine.compare()).
            approvals: Dict mapping category keys to bool (True = use suggested).

        Returns:
            Final CV markdown string.

        Raises:
            AssemblyError: If zero approvals are True.
        """
        if not approvals or not any(approvals.values()):
            raise AssemblyError("At least one category must be approved.")

        parts = []

        for cat_key in CVParser.CATEGORY_ORDER:
            if cat_key == 'professional_experience':
                # Assemble PE header
                pe_entry = comparison.get('professional_experience')
                if pe_entry:
                    approved = approvals.get('professional_experience', False)
                    header = pe_entry['suggested_content'] if approved else pe_entry['original_content']
                    if header:
                        parts.append(header)
                        parts.append('')

                # Assemble PE roles in original sequence
                role_idx = 0
                while True:
                    role_key = f'professional_experience.role_{role_idx}'
                    if role_key not in comparison:
                        break
                    role_entry = comparison[role_key]
                    approved = approvals.get(role_key, False)
                    content = role_entry['suggested_content'] if approved else role_entry['original_content']
                    if content:
                        parts.append(content)
                        parts.append('')
                    role_idx += 1
            else:
                entry = comparison.get(cat_key)
                if not entry:
                    continue
                approved = approvals.get(cat_key, False)
                content = entry['suggested_content'] if approved else entry['original_content']
                if content:
                    parts.append(content)
                    parts.append('')

        # Join and normalize blank lines
        result = '\n'.join(parts)
        result = re.sub(r'\n{3,}', '\n\n', result)
        return result.strip() + '\n'


# ========== APPLICATIONS API ==========

@app.route('/api/applications', methods=['GET'])
def get_applications():
    """Get all applications for a user"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        user_id = request.args.get('userId', 'default-user')
        response = applications_table.query(
            KeyConditionExpression=Key('userId').eq(user_id)
        )
        items = json_safe(response.get('Items', []))
        # Sort by dateApplied descending
        items = sorted(items, key=lambda x: x.get('dateApplied', ''), reverse=True)
        return jsonify(items)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/applications', methods=['POST'])
def create_application():
    """Create or update an application"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        data = request.json
        user_id = data.get('userId', 'default-user')
        
        item = {
            'userId': user_id,
            'id': data.get('id'),
            'company': data.get('company'),
            'role': data.get('role'),
            'jd': data.get('jd', ''),
            'status': data.get('status', 'Applied'),
            'dateApplied': data.get('dateApplied'),
            'followUpDate': data.get('followUpDate', ''),
            'atsScore': data.get('atsScore'),
            'skills': data.get('skills', []),
            'notes': data.get('notes', ''),
            'emailSubject': data.get('emailSubject', ''),
            'emailTo': data.get('emailTo', ''),
            'tailoredCV': data.get('tailoredCV', ''),
            'coverLetter': data.get('coverLetter', ''),
            'emailDraft': data.get('emailDraft', ''),
            'baseCvVersion': data.get('baseCvVersion', ''),
            'timestamp': int(datetime.now(timezone.utc).timestamp())
        }

        # Preserve CV file metadata if present in request data
        if data.get('cvFilename'):
            item['cvFilename'] = data['cvFilename']
        if data.get('cvS3Key'):
            item['cvS3Key'] = data['cvS3Key']
        
        applications_table.put_item(Item=dynamodb_safe(item))
        return jsonify({'status': 'ok', 'id': data.get('id')})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/applications/<app_id>', methods=['PUT'])
def update_application(app_id):
    """Update an application"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        data = request.json
        user_id = data.get('userId', 'default-user')
        
        # Get existing item
        response = applications_table.get_item(
            Key={'userId': user_id, 'id': app_id}
        )
        existing = response.get('Item', {})
        
        # Merge with new data
        item = {**existing, **data}
        item['timestamp'] = int(datetime.now(timezone.utc).timestamp())
        
        applications_table.put_item(Item=dynamodb_safe(item))
        return jsonify({'status': 'ok'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/applications/<app_id>', methods=['DELETE'])
def delete_application(app_id):
    """Delete an application"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        user_id = request.args.get('userId', 'default-user')
        applications_table.delete_item(
            Key={'userId': user_id, 'id': app_id}
        )
        return jsonify({'status': 'ok'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ========== CV UPLOAD API ==========

@app.route('/api/applications/<app_id>/cv', methods=['POST'])
def upload_cv(app_id):
    """Upload a CV file for a specific application"""
    if not S3_AVAILABLE:
        return jsonify({'error': 'Storage service is unavailable'}), 503

    # Extract userId from form data
    user_id = request.form.get('userId')
    if not user_id:
        return jsonify({'error': 'userId is required'}), 400

    # Get the uploaded file
    file = request.files.get('file')

    # Validate the file
    validation_error = validate_cv_file(file)
    if validation_error:
        return jsonify({'error': validation_error}), 400

    # Verify the application exists for this user in DynamoDB
    try:
        response = applications_table.get_item(
            Key={'userId': user_id, 'id': app_id}
        )
    except Exception as e:
        return jsonify({'error': f'Database error: {str(e)}'}), 500

    item = response.get('Item')
    if not item:
        return jsonify({'error': 'Application not found'}), 404

    # Check userId ownership
    if item.get('userId') != user_id:
        return jsonify({'error': 'Insufficient permissions'}), 403

    # Upload file to S3
    filename = os.path.basename(file.filename or '')
    filename = re.sub(r'[^\w\-.]', '_', filename)
    s3_key = f"{user_id}/{app_id}/{filename}"

    try:
        file.seek(0)
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=file.read(),
            ContentType=file.content_type
        )
    except Exception as e:
        return jsonify({'error': 'Failed to upload file to storage'}), 503

    # Update DynamoDB record with CV metadata
    try:
        applications_table.update_item(
            Key={'userId': user_id, 'id': app_id},
            UpdateExpression='SET cvS3Key = :s3key, cvFilename = :fname',
            ExpressionAttributeValues={
                ':s3key': s3_key,
                ':fname': filename
            }
        )
    except Exception as e:
        return jsonify({'error': f'Failed to update application record: {str(e)}'}), 500

    return jsonify({
        'status': 'ok',
        'filename': filename,
        's3Key': s3_key
    }), 200


@app.route('/api/applications/<app_id>/cv', methods=['GET'])
def download_cv(app_id):
    """Download a CV file for a specific application"""
    if not S3_AVAILABLE:
        return jsonify({'error': 'Storage service is unavailable'}), 503

    # Extract userId from query parameter
    user_id = request.args.get('userId')
    if not user_id:
        return jsonify({'error': 'userId is required'}), 400

    # Verify the application exists for this user in DynamoDB
    try:
        response = applications_table.get_item(
            Key={'userId': user_id, 'id': app_id}
        )
    except Exception as e:
        return jsonify({'error': f'Database error: {str(e)}'}), 500

    item = response.get('Item')
    if not item:
        return jsonify({'error': 'Application not found'}), 404

    # Check userId ownership
    if item.get('userId') != user_id:
        return jsonify({'error': 'Insufficient permissions'}), 403

    # Check that CV metadata exists
    s3_key = item.get('cvS3Key')
    filename = item.get('cvFilename')
    if not s3_key or not filename:
        return jsonify({'error': 'No CV associated with this application'}), 404

    # Generate pre-signed URL
    try:
        url = s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': S3_BUCKET, 'Key': s3_key},
            ExpiresIn=900  # 15 minutes
        )
        return jsonify({
            'url': url,
            'filename': filename
        }), 200
    except s3_client.exceptions.NoSuchKey:
        return jsonify({'error': 'CV file not found in storage'}), 404
    except Exception as e:
        if 'NoSuchKey' in str(e):
            return jsonify({'error': 'CV file not found in storage'}), 404
        return jsonify({'error': 'Failed to generate download link'}), 503


@app.route('/api/applications/<app_id>/cv', methods=['DELETE'])
def delete_cv(app_id):
    """Delete a CV file for a specific application"""
    if not S3_AVAILABLE:
        return jsonify({'error': 'Storage service is unavailable'}), 503

    # Extract userId from query parameter
    user_id = request.args.get('userId')
    if not user_id:
        return jsonify({'error': 'userId is required'}), 400

    # Verify the application exists for this user in DynamoDB
    try:
        response = applications_table.get_item(
            Key={'userId': user_id, 'id': app_id}
        )
    except Exception as e:
        return jsonify({'error': f'Database error: {str(e)}'}), 500

    item = response.get('Item')
    if not item:
        return jsonify({'error': 'Application not found'}), 404

    # Check userId ownership
    if item.get('userId') != user_id:
        return jsonify({'error': 'Insufficient permissions'}), 403

    # Check that CV metadata exists
    s3_key = item.get('cvS3Key')
    if not s3_key:
        return jsonify({'error': 'No CV associated with this application'}), 404

    # Delete from S3
    try:
        s3_client.delete_object(Bucket=S3_BUCKET, Key=s3_key)
    except Exception as e:
        return jsonify({'error': 'Failed to delete file from storage'}), 503

    # Remove CV metadata from DynamoDB
    try:
        applications_table.update_item(
            Key={'userId': user_id, 'id': app_id},
            UpdateExpression='REMOVE cvS3Key, cvFilename'
        )
    except Exception as e:
        return jsonify({'error': f'Partial failure: CV deleted from storage but metadata update failed: {str(e)}'}), 500

    return jsonify({'status': 'ok'}), 200


# ========== RUNBOOK API ==========

@app.route('/api/runbook', methods=['GET'])
def get_runbook():
    """Get all runbook questions for a user"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        user_id = request.args.get('userId', 'default-user')
        response = runbook_table.query(
            KeyConditionExpression=Key('userId').eq(user_id)
        )
        items = json_safe(response.get('Items', []))
        return jsonify(items)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/runbook', methods=['POST'])
def create_question():
    """Add an interview question to runbook"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        data = request.json
        user_id = data.get('userId', 'default-user')
        
        item = {
            'userId': user_id,
            'id': data.get('id'),
            'roleKey': data.get('roleKey', ''),
            'question': data.get('question'),
            'type': data.get('type'),
            'round': data.get('round'),
            'outcome': data.get('outcome'),
            'notes': data.get('notes', ''),
            'date': data.get('date'),
            'suggestions': data.get('suggestions'),
            'timestamp': int(datetime.now(timezone.utc).timestamp())
        }
        
        runbook_table.put_item(Item=dynamodb_safe(item))
        return jsonify({'status': 'ok', 'id': data.get('id')})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/runbook/<question_id>', methods=['PUT'])
def update_question(question_id):
    """Update a runbook question"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        data = request.json
        user_id = data.get('userId', 'default-user')
        
        # Get existing item
        response = runbook_table.get_item(
            Key={'userId': user_id, 'id': question_id}
        )
        existing = response.get('Item', {})
        
        # Merge with new data
        item = {**existing, **data}
        item['timestamp'] = int(datetime.now(timezone.utc).timestamp())
        
        runbook_table.put_item(Item=dynamodb_safe(item))
        return jsonify({'status': 'ok'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/runbook/<question_id>', methods=['DELETE'])
def delete_question(question_id):
    """Delete a runbook question"""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'DynamoDB not available'}), 503
    
    try:
        user_id = request.args.get('userId', 'default-user')
        runbook_table.delete_item(
            Key={'userId': user_id, 'id': question_id}
        )
        return jsonify({'status': 'ok'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ========== HEALTH CHECK ==========

@app.route('/api/health', methods=['GET'])
def health():
    """Health check endpoint"""
    try:
        # Test DynamoDB connection
        dynamodb_status = 'disconnected'
        if DYNAMODB_AVAILABLE:
            try:
                applications_table.table_status
                dynamodb_status = 'connected'
            except:
                dynamodb_status = 'disconnected'
        
        return jsonify({
            'status': 'ok',
            'dynamodb': dynamodb_status,
            'docx_support': DOCX_AVAILABLE,
            'llm': 'api-based (OpenAI/Anthropic/Custom)'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'error': str(e)
        }), 500

# ========== LOGIN ==========

@app.route('/api/login', methods=['POST'])
def login():
    if not CAREEROPS_PASSWORD:
        return jsonify({'error': 'Login is not configured'}), 503

    data = request.json or {}
    username = str(data.get('username', ''))
    password = str(data.get('password', ''))

    if (
        hmac.compare_digest(username, CAREEROPS_USERNAME)
        and hmac.compare_digest(password, CAREEROPS_PASSWORD)
    ):
        return jsonify({'ok': True, 'username': username})

    return jsonify({'error': 'Invalid username or password'}), 401

# ========== DOCX SUPPORT ==========

@app.route('/api/extract-docx', methods=['POST'])
def extract_docx():
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'python-docx not installed'}), 400
    try:
        file = request.files.get('file')
        if not file:
            return jsonify({'error': 'No file uploaded'}), 400
        
        from docx import Document
        doc = Document(file)
        text = '\n'.join([para.text for para in doc.paragraphs])
        
        return jsonify({'text': text})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/generate-docx', methods=['POST'])
def generate_docx():
    """Convert markdown CV text to a downloadable DOCX file."""
    if not DOCX_AVAILABLE:
        return jsonify({'error': 'python-docx not installed. Run: pip install python-docx'}), 503

    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    data = request.json
    cv_text = (data or {}).get('cv_text', '')
    if not cv_text.strip():
        return jsonify({'error': 'cv_text is required'}), 400

    try:
        doc = Document()

        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10.5)

        # Set narrow margins for ATS-friendly single-page layout
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        lines = cv_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # H1: Name / title (# ...)
            if stripped.startswith('# '):
                text = stripped[2:].strip()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(16)
                p.space_after = Pt(2)
                i += 1
                continue

            # H2: Section headers (## ...)
            if stripped.startswith('## '):
                text = stripped[3:].strip()
                p = doc.add_paragraph()
                run = p.add_run(text.upper())
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                p.space_before = Pt(10)
                p.space_after = Pt(3)
                # Add a bottom border effect via a thin line
                i += 1
                continue

            # H3: Sub-headers / role titles (### ...)
            if stripped.startswith('### '):
                text = stripped[4:].strip()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(10.5)
                p.space_before = Pt(6)
                p.space_after = Pt(1)
                i += 1
                continue

            # Bullet points (- ... or * ...)
            if stripped.startswith('- ') or stripped.startswith('* '):
                text = stripped[2:].strip()
                p = doc.add_paragraph(style='List Bullet')
                # Handle bold markers **...**
                parts = re.split(r'\*\*(.*?)\*\*', text)
                for j, part in enumerate(parts):
                    if j % 2 == 1:
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(0)
                i += 1
                continue

            # Regular paragraph (handle inline **bold** and contact info lines)
            p = doc.add_paragraph()
            # Check if it's a contact/info line (contains | separators)
            if '|' in stripped and len(stripped.split('|')) >= 3:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            else:
                parts = re.split(r'\*\*(.*?)\*\*', stripped)
                for j, part in enumerate(parts):
                    if j % 2 == 1:
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            p.space_after = Pt(2)
            i += 1

        # Write to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = (data or {}).get('filename', 'CV.docx')
        if not filename.lower().endswith('.docx'):
            filename += '.docx'

        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except Exception as e:
        return jsonify({'error': f'DOCX generation failed: {str(e)}'}), 500

# ========== LLM PROVIDER CONFIGURATION (Secure) ==========

@app.route('/api/llm-providers', methods=['GET'])
def list_llm_providers():
    """List all LLM provider configs for a user. API keys are masked."""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'Database not available'}), 503

    user_id = request.args.get('userId')
    if not user_id:
        return jsonify({'error': 'userId is required'}), 400

    try:
        response = llm_providers_table.query(
            KeyConditionExpression=Key('userId').eq(user_id)
        )
        items = json_safe(response.get('Items', []))

        # Mask API keys before returning
        for item in items:
            if item.get('apiKeyEncrypted'):
                try:
                    raw_key = decrypt_value(item['apiKeyEncrypted'])
                    item['apiKeyMasked'] = mask_key(raw_key)
                except Exception:
                    item['apiKeyMasked'] = '••••••••'
                del item['apiKeyEncrypted']
            else:
                item['apiKeyMasked'] = ''
            item.pop('apiKeyEncrypted', None)

        items = sorted(items, key=lambda x: x.get('updatedAt', ''), reverse=True)
        return jsonify(items)
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/llm-providers', methods=['POST'])
def create_llm_provider():
    """Create or update an LLM provider config. API key is encrypted at rest."""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'Database not available'}), 503
    if not fernet:
        return jsonify({'error': 'Encryption not available. Set CAREEROPS_ENCRYPTION_KEY env var.'}), 503

    data = request.json
    user_id = data.get('userId')
    if not user_id:
        return jsonify({'error': 'userId is required'}), 400

    provider_type = data.get('provider', '').strip()
    if not provider_type:
        return jsonify({'error': 'provider is required (openai, anthropic, custom)'}), 400

    config_id = data.get('id') or f"{provider_type}-{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}"
    api_key = data.get('apiKey', '')
    model = data.get('model', '')
    base_url = data.get('baseUrl', '')
    label = data.get('label', '') or f"{provider_type} — {model or 'default'}"

    item = {
        'userId': user_id,
        'id': config_id,
        'provider': provider_type,
        'model': model,
        'baseUrl': base_url,
        'label': label,
        'isActive': data.get('isActive', False),
        'updatedAt': datetime.now(timezone.utc).isoformat(),
    }

    # Encrypt API key if provided
    if api_key:
        item['apiKeyEncrypted'] = encrypt_value(api_key)

    # If updating an existing provider and no new key provided, preserve existing encrypted key
    if not api_key and data.get('id'):
        try:
            existing = llm_providers_table.get_item(Key={'userId': user_id, 'id': config_id})
            existing_item = existing.get('Item')
            if existing_item and existing_item.get('apiKeyEncrypted'):
                item['apiKeyEncrypted'] = existing_item['apiKeyEncrypted']
        except Exception:
            pass

    try:
        llm_providers_table.put_item(Item=dynamodb_safe(item))

        # If isActive, deactivate others
        if data.get('isActive'):
            response = llm_providers_table.query(
                KeyConditionExpression=Key('userId').eq(user_id)
            )
            for other in response.get('Items', []):
                if other['id'] != config_id and other.get('isActive'):
                    llm_providers_table.update_item(
                        Key={'userId': user_id, 'id': other['id']},
                        UpdateExpression='SET isActive = :f',
                        ExpressionAttributeValues={':f': False}
                    )

        return jsonify({
            'status': 'ok',
            'id': config_id,
            'apiKeyMasked': mask_key(api_key) if api_key else ''
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/llm-providers/<config_id>', methods=['DELETE'])
def delete_llm_provider(config_id):
    """Delete an LLM provider config."""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'Database not available'}), 503

    user_id = request.args.get('userId')
    if not user_id:
        return jsonify({'error': 'userId is required'}), 400

    try:
        llm_providers_table.delete_item(Key={'userId': user_id, 'id': config_id})
        return jsonify({'status': 'ok'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/llm-providers/<config_id>/activate', methods=['POST'])
def activate_llm_provider(config_id):
    """Set a provider config as the active one (deactivates others)."""
    if not DYNAMODB_AVAILABLE:
        return jsonify({'error': 'Database not available'}), 503

    data = request.json or {}
    user_id = data.get('userId') or request.args.get('userId')
    if not user_id:
        return jsonify({'error': 'userId is required'}), 400

    try:
        # Verify it exists
        existing = llm_providers_table.get_item(Key={'userId': user_id, 'id': config_id})
        if not existing.get('Item'):
            return jsonify({'error': 'Provider config not found'}), 404

        # Deactivate all others
        response = llm_providers_table.query(
            KeyConditionExpression=Key('userId').eq(user_id)
        )
        for item in response.get('Items', []):
            if item['id'] != config_id and item.get('isActive'):
                llm_providers_table.update_item(
                    Key={'userId': user_id, 'id': item['id']},
                    UpdateExpression='SET isActive = :f',
                    ExpressionAttributeValues={':f': False}
                )

        # Activate this one
        llm_providers_table.update_item(
            Key={'userId': user_id, 'id': config_id},
            UpdateExpression='SET isActive = :t',
            ExpressionAttributeValues={':t': True}
        )

        return jsonify({'status': 'ok'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def get_active_provider_key(user_id):
    """Fetch the decrypted API key for the active provider. Returns (key, provider_config) or (None, None)."""
    if not DYNAMODB_AVAILABLE or not fernet:
        print(f"[get_active_provider_key] DYNAMODB_AVAILABLE={DYNAMODB_AVAILABLE}, fernet={'set' if fernet else 'None'}")
        return None, None
    try:
        response = llm_providers_table.query(
            KeyConditionExpression=Key('userId').eq(user_id)
        )
        items = response.get('Items', [])
        print(f"[get_active_provider_key] Found {len(items)} providers for userId={user_id}")
        for item in items:
            print(f"  - id={item.get('id')} provider={item.get('provider')} isActive={item.get('isActive')} hasKey={'yes' if item.get('apiKeyEncrypted') else 'no'}")
            if item.get('isActive') and item.get('apiKeyEncrypted'):
                try:
                    key = decrypt_value(item['apiKeyEncrypted'])
                    print(f"  → Decrypted key successfully (length={len(key)})")
                    return key, json_safe(item)
                except Exception as dec_err:
                    print(f"  → DECRYPTION FAILED: {dec_err} — key was likely saved with a different CAREEROPS_ENCRYPTION_KEY")
                    return None, None
        print(f"[get_active_provider_key] No active provider with encrypted key found")
        return None, None
    except Exception as e:
        print(f"[get_active_provider_key] Exception: {e}")
        return None, None


def get_provider_key_by_id(user_id, config_id):
    """Fetch the decrypted API key for a specific provider config."""
    if not DYNAMODB_AVAILABLE or not fernet:
        return None, None
    try:
        response = llm_providers_table.get_item(Key={'userId': user_id, 'id': config_id})
        item = response.get('Item')
        if item and item.get('apiKeyEncrypted'):
            key = decrypt_value(item['apiKeyEncrypted'])
            return key, json_safe(item)
        return None, json_safe(item) if item else None
    except Exception:
        return None, None


# ========== CV COMPARISON API ==========


def save_comparison_state(user_id, app_id, comparison_data, approvals, scores):
    """Persist comparison state to the application record in DynamoDB.

    Args:
        user_id: User identifier.
        app_id: Application identifier.
        comparison_data: Comparison dict from ComparisonEngine.
        approvals: Dict mapping category keys to bool.
        scores: Dict with baseline_score, max_score, etc.
    """
    if not DYNAMODB_AVAILABLE:
        return

    try:
        response = applications_table.get_item(
            Key={'userId': user_id, 'id': app_id}
        )
        item = response.get('Item') or {'userId': user_id, 'id': app_id}

        item['comparisonState'] = {
            'comparison': comparison_data,
            'approvals': approvals,
            'scores': scores,
            'updatedAt': datetime.now(timezone.utc).isoformat(),
        }
        item['timestamp'] = int(datetime.now(timezone.utc).timestamp())

        applications_table.put_item(Item=dynamodb_safe(item))
    except Exception as e:
        print(f"[save_comparison_state] Error: {e}")


def load_comparison_state(user_id, app_id):
    """Load comparison state from the application record in DynamoDB.

    Args:
        user_id: User identifier.
        app_id: Application identifier.

    Returns:
        Comparison state dict or None if not found.
    """
    if not DYNAMODB_AVAILABLE:
        return None

    try:
        response = applications_table.get_item(
            Key={'userId': user_id, 'id': app_id}
        )
        item = response.get('Item')
        if item and item.get('comparisonState'):
            return json_safe(item['comparisonState'])
        return None
    except Exception as e:
        print(f"[load_comparison_state] Error: {e}")
        return None


@app.route('/api/compare-cv', methods=['POST'])
def compare_cv():
    """Compare original and tailored CVs, returning category-level diff.

    Expects JSON body:
        userId: User identifier.
        appId: Application identifier.
        originalCv: Original CV markdown string.
        tailoredCv: Tailored CV markdown string.
        baselineScore: Baseline ATS score (int).
        categoryScores: Dict mapping category keys to score impact (int).

    Returns:
        JSON with comparison, baseline_score, max_score, total_changes.
    """
    try:
        data = request.json
        user_id = data.get('userId', 'default-user')
        app_id = data.get('appId', '')
        original_cv = data.get('originalCv', '')
        tailored_cv = data.get('tailoredCv', '')
        baseline_score = data.get('baselineScore', 0)
        category_scores = data.get('categoryScores', {})

        parser = CVParser()

        # Parse both CVs
        try:
            original_parsed = parser.parse(original_cv)
        except ParseError as e:
            return jsonify({'error': f'Failed to parse original CV: {str(e)}'}), 400

        try:
            suggested_parsed = parser.parse(tailored_cv)
        except ParseError as e:
            return jsonify({'error': f'Failed to parse tailored CV: {str(e)}'}), 400

        # Compare
        engine = ComparisonEngine()
        try:
            comparison = engine.compare(original_parsed, suggested_parsed)
        except ComparisonError as e:
            return jsonify({'error': f'Comparison failed: {str(e)}'}), 400

        # Attach score_impact from categoryScores
        for key, entry in comparison.items():
            if key in category_scores:
                entry['score_impact'] = category_scores[key]

        # Calculate totals
        total_changes = sum(1 for entry in comparison.values() if entry['changed'])
        max_score = max(0, min(100, baseline_score + sum(category_scores.values())))

        # Persist state if appId provided
        if app_id:
            approvals_to_save = data.get('_approvals', {}) if data.get('_saveState') else {}
            scores = {
                'baseline_score': baseline_score,
                'max_score': max_score,
                'category_scores': category_scores,
            }
            save_comparison_state(user_id, app_id, comparison, approvals_to_save, scores)

        return jsonify({
            'comparison': comparison,
            'baseline_score': baseline_score,
            'max_score': max_score,
            'total_changes': total_changes,
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/assemble-cv', methods=['POST'])
def assemble_cv():
    """Assemble final CV from comparison data and user approvals.

    Expects JSON body:
        userId: User identifier.
        appId: Application identifier.
        comparison: Comparison dict (from /api/compare-cv).
        approvals: Dict mapping category keys to bool.

    Returns:
        JSON with assembled_cv and final_score.
    """
    try:
        data = request.json
        user_id = data.get('userId', 'default-user')
        app_id = data.get('appId', '')
        comparison = data.get('comparison', {})
        approvals = data.get('approvals', {})

        assembler = FinalAssembler()

        try:
            assembled_cv = assembler.assemble(comparison, approvals)
        except AssemblyError as e:
            return jsonify({'error': str(e)}), 400

        # Compute final ATS score
        # Build impacts and approval dicts from comparison
        impacts = {}
        for key, entry in comparison.items():
            impacts[key] = entry.get('score_impact', 0)

        # Use baseline from request body, fall back to persisted state
        baseline_score = data.get('baselineScore', 0)
        if not baseline_score and app_id:
            state = load_comparison_state(user_id, app_id)
            if state and isinstance(state, dict) and state.get('scores'):
                baseline_score = state['scores'].get('baseline_score', 0)

        score_result = compute_ats_score(baseline_score, impacts, approvals)
        final_score = score_result['current_score']

        # Persist updated state
        if app_id:
            scores = {
                'baseline_score': baseline_score,
                'max_score': score_result['max_score'],
                'category_scores': impacts,
            }
            save_comparison_state(user_id, app_id, comparison, approvals, scores)

        return jsonify({
            'assembled_cv': assembled_cv,
            'final_score': final_score,
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ========== LLM GENERATION ==========

@app.route('/api/generate', methods=['POST'])
def generate():
    try:
        data = request.json
        provider = data.get('provider', 'openai')
        system = data.get('system', '')
        user_msg = data.get('user_message', '')
        model = data.get('model', 'gpt-4o-mini')
        user_id = data.get('userId', '')
        provider_id = data.get('providerId', '')

        # Fetch API key server-side from encrypted storage
        api_key = None
        provider_config = None

        print(f"[generate] provider={provider}, userId={user_id}, providerId={provider_id}")

        # Try to fetch key from DynamoDB by provider ID
        if provider_id and user_id:
            api_key, provider_config = get_provider_key_by_id(user_id, provider_id)
        elif user_id:
            api_key, provider_config = get_active_provider_key(user_id)

        # Fallback to environment variable for OpenAI
        if not api_key and provider == 'openai':
            api_key = OPENAI_API_KEY
            if api_key:
                print(f"[generate] Using OPENAI_API_KEY env var fallback")

        print(f"[generate] api_key resolved: {'yes' if api_key else 'NO'}")

        if provider == 'openai':
            return generate_openai(data, system, user_msg, model, api_key)
        elif provider == 'custom':
            base_url = (provider_config or {}).get('baseUrl', '') if provider_config else ''
            return generate_custom(data, system, user_msg, model, api_key, base_url)
        else:
            # Default to OpenAI for any unrecognized provider
            return generate_openai(data, system, user_msg, model, api_key)

    except requests.exceptions.Timeout:
        return jsonify({'error': 'Request timed out. Try again.'}), 504
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def generate_openai(data, system, user_msg, model, api_key=None):
    """Handle OpenAI ChatGPT API calls. API key fetched server-side."""
    if not api_key:
        api_key = OPENAI_API_KEY
    if not api_key:
        return jsonify({'error': 'OpenAI API key not configured. Delete and re-add the provider in Settings (encryption key may have changed).'}), 400

    max_tokens = data.get('max_tokens', 12000)
    messages = []
    if system:
        messages.append({'role': 'system', 'content': system})
    messages.append({'role': 'user', 'content': user_msg})

    try:
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers={
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {api_key}'
            },
            json={
                'model': model,
                'messages': messages,
                'temperature': 0.7,
                'max_tokens': max_tokens,
                'stream': False
            },
            timeout=180
        )

        if response.status_code == 401:
            return jsonify({'error': 'Invalid OpenAI API key. Check your key in Settings.'}), 401
        if response.status_code == 429:
            return jsonify({'error': 'OpenAI rate limit exceeded. Wait a moment and try again.'}), 429
        if response.status_code == 404:
            return jsonify({'error': f'Model "{model}" not found. Check your OpenAI plan supports this model.'}), 404
        if response.status_code != 200:
            detail = response.text
            try:
                detail = response.json().get('error', {}).get('message', response.text)
            except Exception:
                pass
            return jsonify({'error': f'OpenAI error ({response.status_code}): {detail}'}), response.status_code

        result = response.json()
        message_content = result['choices'][0]['message']['content']

        def fix_json_strings(s):
            def escape_inner(m):
                inner = m.group(1)
                inner = inner.replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
                return '"' + inner + '"'
            return re.sub(r'"((?:[^"\\]|\\.)*)"', escape_inner, s, flags=re.DOTALL)

        stripped = message_content.strip()
        if stripped.startswith('```'):
            stripped = re.sub(r'^```(?:json)?\s*', '', stripped)
            stripped = re.sub(r'```\s*$', '', stripped).strip()
        try:
            json.loads(stripped)
        except (json.JSONDecodeError, ValueError):
            stripped = fix_json_strings(stripped)

        # Build response, passing through score data if present in LLM output
        response_data = {
            'content': [{'type': 'text', 'text': stripped}]
        }
        try:
            parsed_llm = json.loads(stripped)
            if isinstance(parsed_llm, dict):
                if 'baseline_score' in parsed_llm:
                    response_data['baseline_score'] = parsed_llm['baseline_score']
                if 'category_scores' in parsed_llm:
                    response_data['category_scores'] = parsed_llm['category_scores']
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

        return jsonify(response_data)

    except requests.exceptions.Timeout:
        return jsonify({'error': 'OpenAI request timed out. Try again.'}), 504
    except requests.exceptions.ConnectionError:
        return jsonify({'error': 'Cannot reach OpenAI API. Check your network connection.'}), 503
    except Exception as e:
        return jsonify({'error': f'OpenAI request failed: {str(e)}'}), 500


def generate_custom(data, system, user_msg, model, api_key=None, base_url=''):
    """Handle custom OpenAI-compatible API calls."""
    if not base_url:
        return jsonify({'error': 'Custom provider base URL not configured. Update it in Settings.'}), 400

    max_tokens = data.get('max_tokens', 12000)
    messages = []
    if system:
        messages.append({'role': 'system', 'content': system})
    messages.append({'role': 'user', 'content': user_msg})

    headers = {'Content-Type': 'application/json'}
    if api_key:
        headers['Authorization'] = f'Bearer {api_key}'

    # Ensure base_url ends with /v1/chat/completions or similar
    url = base_url.rstrip('/')
    if not url.endswith('/chat/completions'):
        url = url + '/chat/completions'

    try:
        response = requests.post(
            url,
            headers=headers,
            json={
                'model': model,
                'messages': messages,
                'temperature': 0.7,
                'max_tokens': max_tokens,
                'stream': False
            },
            timeout=180
        )

        if response.status_code != 200:
            detail = response.text
            try:
                detail = response.json().get('error', {}).get('message', response.text)
            except Exception:
                pass
            return jsonify({'error': f'Custom provider error ({response.status_code}): {detail}'}), response.status_code

        result = response.json()
        message_content = result['choices'][0]['message']['content']

        def fix_json_strings(s):
            def escape_inner(m):
                inner = m.group(1)
                inner = inner.replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
                return '"' + inner + '"'
            return re.sub(r'"((?:[^"\\]|\\.)*)"', escape_inner, s, flags=re.DOTALL)

        stripped = message_content.strip()
        if stripped.startswith('```'):
            stripped = re.sub(r'^```(?:json)?\s*', '', stripped)
            stripped = re.sub(r'```\s*$', '', stripped).strip()
        try:
            json.loads(stripped)
        except (json.JSONDecodeError, ValueError):
            stripped = fix_json_strings(stripped)

        # Build response, passing through score data if present in LLM output
        response_data = {
            'content': [{'type': 'text', 'text': stripped}]
        }
        try:
            parsed_llm = json.loads(stripped)
            if isinstance(parsed_llm, dict):
                if 'baseline_score' in parsed_llm:
                    response_data['baseline_score'] = parsed_llm['baseline_score']
                if 'category_scores' in parsed_llm:
                    response_data['category_scores'] = parsed_llm['category_scores']
        except (json.JSONDecodeError, ValueError, TypeError):
            pass

        return jsonify(response_data)

    except requests.exceptions.Timeout:
        return jsonify({'error': 'Custom provider request timed out.'}), 504
    except requests.exceptions.ConnectionError:
        return jsonify({'error': f'Cannot reach custom provider at {base_url}. Check the URL.'}), 503
    except Exception as e:
        return jsonify({'error': f'Custom provider request failed: {str(e)}'}), 500


if __name__ == '__main__':
    port = int(os.environ.get('FLASK_PORT', 5001))

    if DYNAMODB_AVAILABLE:
        print(f"🔍 Checking DynamoDB...")
        try:
            applications_table.table_status
            print("✅ DynamoDB connected!")
        except Exception as e:
            print(f"⚠️  DynamoDB connection error: {e}")
    else:
        print("\n⚠️  DynamoDB not available")
    
    print(f"🔑 Encryption: {'configured' if fernet else 'NOT available'}")
    print(f"🚀 Backend running on http://localhost:{port}")
    print(f"📄 Frontend: http://localhost:8000/careerops.html")
    print(f"✓ Health check: http://localhost:{port}/api/health")
    print(f"🤖 LLM: API-based (OpenAI/Anthropic/Custom via Settings)")
    host = os.environ.get('FLASK_HOST', '127.0.0.1')
    debug = os.environ.get('FLASK_DEBUG', 'false').lower() in ('1', 'true', 'yes')
    app.run(host=host, port=port, debug=debug)
